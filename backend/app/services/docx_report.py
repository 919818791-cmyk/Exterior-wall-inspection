from __future__ import annotations

from collections import Counter
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from re import IGNORECASE, Match, search
from typing import Any, Callable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, RGBColor
from docx.text.run import Run
from PIL import Image, ImageOps, UnidentifiedImageError

from app.services.defect_annotation import (
    DefectAnnotationError,
    defect_label,
    draw_defect_annotations,
    finite_number,
)
from app.services.defect_numbering import number_defects


BACKEND_ROOT = Path(__file__).resolve().parents[2]
REPORT_TEMPLATE_DIR = BACKEND_ROOT / "templates" / "reports"
FORMAL_TEMPLATE = REPORT_TEMPLATE_DIR / "正式报告示例.docx"

TABLE_IMAGE_WIDTH_EMU = 1_440_000
TABLE_IMAGE_HEIGHT_EMU = 1_080_000
TABLE_IMAGE_MAX_WIDTH = 1_200
TABLE_IMAGE_MAX_HEIGHT = 900
TABLE_IMAGE_JPEG_QUALITY = 82
REPORT_TITLE_SUFFIX = "外立面表观病害筛查简报"
REPORT_INTRO_TEXT = (
    "经对巡检结果进行空间定位与尺度估算，得到以下疑似病害位置。"
    "深度估计结果存在模型与相机参数误差，建议结合现场复核。"
)

ObjectReader = Callable[[str, str], bytes]


class DocxReportExportError(RuntimeError):
    pass


def _text(value: Any, *, fallback: str = "-") -> str:
    if value is None or value == "":
        return fallback
    return str(value)


def _photo_key(item: dict[str, Any]) -> str:
    if item.get("photo_id"):
        return f"photo:{item['photo_id']}"
    if item.get("id"):
        return f"photo:{item['id']}"
    filename = item.get("photo_filename") or item.get("original_filename")
    return f"filename:{filename}" if filename else ""


def _filename_variant_match(filename: Any) -> Match[str] | None:
    return search(r"_([vt])(?:\.[^.]+)$", str(filename or "").strip(), flags=IGNORECASE)


def _filename_variant(filename: Any) -> str | None:
    match = _filename_variant_match(filename)
    if match is None:
        return None
    return "thermal" if match.group(1).lower() == "t" else "visible"


def _photo_pair_key(filename: Any) -> str | None:
    match = search(r"^(.*)_([vt])(?:\.[^.]+)$", str(filename or "").strip(), flags=IGNORECASE)
    return match.group(1).casefold() if match else None


def _is_thermal_photo(photo: dict[str, Any]) -> bool:
    named_variant = _filename_variant(photo.get("original_filename"))
    if named_variant:
        return named_variant == "thermal"
    metadata = photo.get("metadata_json") or {}
    return (
        photo.get("photo_type") == "thermal"
        or photo.get("thermal_imaging_available") is True
        or metadata.get("thermal_imaging_available") is True
    )


def _photo_rows(data: dict[str, Any]) -> list[dict[str, Any]]:
    photos = [deepcopy(photo) for photo in data.get("photos") or [] if isinstance(photo, dict)]
    photo_key_by_id = {
        str(photo["id"]): _photo_key(photo)
        for photo in photos
        if photo.get("id")
    }
    photo_key_by_filename = {
        str(photo["original_filename"]): _photo_key(photo)
        for photo in photos
        if photo.get("original_filename")
    }
    defects_by_photo: dict[str, list[dict[str, Any]]] = {}
    orphan_defects: list[dict[str, Any]] = []
    for defect in data.get("defects") or []:
        if not isinstance(defect, dict):
            continue
        key = (
            photo_key_by_id.get(str(defect.get("photo_id")))
            or photo_key_by_filename.get(str(defect.get("photo_filename")))
            or _photo_key(defect)
        )
        if key:
            defects_by_photo.setdefault(key, []).append(defect)
        else:
            orphan_defects.append(defect)

    rows: list[dict[str, Any]] = []
    consumed_keys: set[str] = set()
    for index, photo in enumerate(photos):
        key = _photo_key(photo) or f"photo-index:{index}"
        consumed_keys.add(key)
        rows.append({
            "key": key,
            "photo": photo,
            "defects": defects_by_photo.get(key, []),
        })

    for key, defects in defects_by_photo.items():
        if key in consumed_keys or not defects:
            continue
        filename = defects[0].get("photo_filename")
        rows.append({
            "key": key,
            "photo": {
                "id": defects[0].get("photo_id"),
                "original_filename": filename,
                "photo_type": _filename_variant(filename) or "visible",
            },
            "defects": defects,
        })

    if orphan_defects:
        rows.append({
            "key": "orphan-defects",
            "photo": {"original_filename": orphan_defects[0].get("photo_filename")},
            "defects": orphan_defects,
        })
    return rows


def _paired_photo_rows(data: dict[str, Any]) -> list[dict[str, Any]]:
    rows = _photo_rows(data)
    consumed_indexes: set[int] = set()
    pairs: list[dict[str, Any]] = []

    for index, row in enumerate(rows):
        if index in consumed_indexes:
            continue
        consumed_indexes.add(index)

        photo = row["photo"]
        filename = photo.get("original_filename")
        named_variant = _filename_variant(filename)
        pair_key = _photo_pair_key(filename)
        matched_index = -1
        if named_variant and pair_key:
            for candidate_index, candidate in enumerate(rows):
                if candidate_index == index or candidate_index in consumed_indexes:
                    continue
                candidate_filename = candidate["photo"].get("original_filename")
                if (
                    _photo_pair_key(candidate_filename) == pair_key
                    and _filename_variant(candidate_filename) not in (None, named_variant)
                ):
                    matched_index = candidate_index
                    break

        matched = rows[matched_index] if matched_index >= 0 else None
        if matched_index >= 0:
            consumed_indexes.add(matched_index)
        row_variant = named_variant or ("thermal" if _is_thermal_photo(photo) else "visible")
        pairs.append({
            "visible": row if row_variant == "visible" else matched,
            "thermal": row if row_variant == "thermal" else matched,
        })

    return pairs


def _defect_description(defects: list[dict[str, Any]]) -> str:
    labels = [
        defect_label(defect.get("defect_type"))
        for defect in defects
    ]
    counts = Counter(labels)
    if not counts:
        return "未检出明显缺陷"
    return "\n".join(
        f"疑似{label}: {count}处"
        for label, count in counts.items()
    )


def _metadata_text(photo: dict[str, Any] | None) -> str:
    if not photo:
        return "未知立面\n--"
    orientation = _text(photo.get("facade_orientation"), fallback="未知立面")
    altitude = photo.get("relative_altitude")
    try:
        altitude_text = f"{float(altitude):.1f} m" if altitude not in (None, "") else "--"
    except (TypeError, ValueError):
        altitude_text = "--"
    return f"{orientation}\n{altitude_text}"


def _defect_measurement_text(value: Any) -> str | None:
    measurement = finite_number(value)
    if measurement is None or measurement < 0:
        return None
    return f"{measurement:.3f}"


def _append_prototype_run(paragraph, prototype: Any | None, value: str) -> Run:
    if prototype is None:
        return paragraph.add_run(value)
    run_element = deepcopy(prototype)
    paragraph._p.append(run_element)
    run = Run(run_element, paragraph)
    run.text = value
    return run


def _set_defect_details(
    paragraph,
    defects: list[dict[str, Any]],
    run_prototypes: list[Any],
) -> None:
    _clear_paragraph_content(paragraph)
    prototype = lambda index: (
        run_prototypes[min(index, len(run_prototypes) - 1)]
        if run_prototypes
        else None
    )
    if not defects:
        _append_prototype_run(paragraph, prototype(0), "—")
        return

    detail_items: list[tuple[dict[str, Any], bool, str] | None] = []
    has_missing_parameters = False
    for defect in defects:
        is_crack = defect.get("defect_type") == "crack"
        measurement_text = _defect_measurement_text(
            defect.get("length") if is_crack else defect.get("area")
        )
        if measurement_text is None:
            has_missing_parameters = True
        else:
            detail_items.append((defect, is_crack, measurement_text))
    if has_missing_parameters:
        detail_items.append(None)

    for index, detail_item in enumerate(detail_items):
        if detail_item is None:
            last_run = _append_prototype_run(paragraph, prototype(0), "参数不足")
            if index < len(detail_items) - 1:
                last_run.add_break()
            continue

        defect, is_crack, measurement_text = detail_item
        defect_no = _text(defect.get("defect_no"), fallback="缺陷")
        _append_prototype_run(paragraph, prototype(0), defect_no)
        estimated = defect.get(
            "length_estimated" if is_crack else "area_estimated"
        ) is True
        _append_prototype_run(paragraph, prototype(1), "≈" if estimated else " ")
        _append_prototype_run(paragraph, prototype(2), measurement_text)
        last_run = _append_prototype_run(
            paragraph,
            prototype(3),
            "m" if is_crack else "m²",
        )
        if index < len(detail_items) - 1:
            last_run.add_break()


def _compressed_table_image(
    original_bytes: bytes,
    defects: list[dict[str, Any]] | None = None,
) -> BytesIO:
    try:
        with Image.open(BytesIO(original_bytes)) as source:
            image = ImageOps.exif_transpose(source)
            if image.mode in ("RGBA", "LA") or (image.mode == "P" and "transparency" in image.info):
                rgba = image.convert("RGBA")
                flattened = Image.new("RGB", rgba.size, "white")
                flattened.paste(rgba, mask=rgba.getchannel("A"))
                image = flattened
            else:
                image = image.convert("RGB")

            draw_defect_annotations(
                image,
                defects or [],
                numbered_labels=True,
                line_width_ratio=0.006,
                min_line_width=5,
                font_size_ratio=0.055,
                min_font_size=32,
                max_font_size=96,
            )

            available_width = min(
                image.width,
                (image.height * 4) // 3,
                TABLE_IMAGE_MAX_WIDTH,
            )
            target_width = max(4, (available_width // 4) * 4)
            target_height = target_width * 3 // 4
            if target_height > TABLE_IMAGE_MAX_HEIGHT:
                target_height = TABLE_IMAGE_MAX_HEIGHT
                target_width = target_height * 4 // 3
            image = ImageOps.fit(
                image,
                (target_width, target_height),
                method=Image.Resampling.LANCZOS,
                centering=(0.5, 0.5),
            )
            output = BytesIO()
            image.save(
                output,
                format="JPEG",
                quality=TABLE_IMAGE_JPEG_QUALITY,
                optimize=True,
                progressive=True,
                subsampling="4:2:0",
            )
            output.seek(0)
            return output
    except DefectAnnotationError as exc:
        raise DocxReportExportError(f"{exc} DOCX 导出已终止。") from exc
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise DocxReportExportError("检测照片无法读取或格式不受支持，DOCX 导出已终止。") from exc


def _clear_paragraph_content(paragraph) -> None:
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)


def _replace_cell_paragraphs(cell, paragraph_prototypes: list[Any]) -> None:
    cell_element = cell._tc
    for child in list(cell_element):
        if child.tag != qn("w:tcPr"):
            cell_element.remove(child)
    for prototype in paragraph_prototypes:
        cell_element.append(deepcopy(prototype))


def _set_paragraph_text(paragraph, value: str, *, blue: bool = False) -> None:
    runs = paragraph.runs
    if runs:
        run = runs[0]
        run.text = value
        for extra_run in runs[1:]:
            extra_run._element.getparent().remove(extra_run._element)
    else:
        run = paragraph.add_run(value)
    if blue:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)


def _read_photo_image(
    photo: dict[str, Any],
    defects: list[dict[str, Any]],
    read_object: ObjectReader | None,
) -> BytesIO | None:
    bucket = photo.get("storage_bucket")
    object_key = photo.get("storage_object_key")
    if read_object is None or not bucket or not object_key:
        return None
    try:
        original_bytes = read_object(str(bucket), str(object_key))
    except Exception as exc:
        filename = _text(photo.get("original_filename"), fallback="检测照片")
        raise DocxReportExportError(f"{filename} 原图读取失败，DOCX 导出已终止。") from exc
    return _compressed_table_image(original_bytes, defects)


def _set_photo_cell(
    cell,
    row: dict[str, Any] | None,
    *,
    read_object: ObjectReader | None,
    populated_prototypes: list[Any],
    empty_prototypes: list[Any],
) -> None:
    if row is None:
        _replace_cell_paragraphs(cell, empty_prototypes)
        return

    _replace_cell_paragraphs(cell, populated_prototypes)
    photo = row["photo"]
    image_stream = _read_photo_image(photo, row["defects"], read_object)
    image_paragraph = cell.paragraphs[0]
    _clear_paragraph_content(image_paragraph)
    if image_stream is not None:
        run = image_paragraph.add_run()
        run.add_picture(
            image_stream,
            width=Emu(TABLE_IMAGE_WIDTH_EMU),
            height=Emu(TABLE_IMAGE_HEIGHT_EMU),
        )
    filename = _text(photo.get("original_filename"), fallback="检测结果照片")
    _set_paragraph_text(cell.paragraphs[1], filename)


def _set_title(document, data: dict[str, Any], report_title: str, report_no: str) -> None:
    if len(document.paragraphs) < 3:
        raise DocxReportExportError("DOCX 模板缺少标题或说明段落。")
    project = data.get("project") or {}
    detection_name = str(project.get("name") or report_title or report_no).strip()
    title = f"{detection_name}-{REPORT_TITLE_SUFFIX}" if detection_name else REPORT_TITLE_SUFFIX
    _set_paragraph_text(document.paragraphs[1], title)
    _set_paragraph_text(document.paragraphs[2], REPORT_INTRO_TEXT)
    document.core_properties.title = title
    document.core_properties.subject = report_no


def _remove_trailing_empty_paragraph(document) -> None:
    if not document.paragraphs or document.paragraphs[-1].text:
        return
    paragraph_element = document.paragraphs[-1]._p
    body = document._body._element
    if paragraph_element.getnext() is body.sectPr:
        body.remove(paragraph_element)


def _populate_result_table(
    document,
    data: dict[str, Any],
    read_object: ObjectReader | None,
) -> None:
    if not document.tables or len(document.tables[0].rows) < 2:
        raise DocxReportExportError("DOCX 模板缺少结果表格数据行原型。")
    table = document.tables[0]
    if len(table.rows[0].cells) < 6 or len(table.rows[1].cells) < 6:
        raise DocxReportExportError("DOCX 模板缺少“缺陷详情”列。")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if header_properties.find(qn("w:tblHeader")) is None:
        header_properties.append(OxmlElement("w:tblHeader"))
    prototype_row = deepcopy(table.rows[1]._tr)
    populated_photo_prototypes = [deepcopy(paragraph._p) for paragraph in table.rows[1].cells[1].paragraphs]
    empty_photo_prototypes = [deepcopy(paragraph._p) for paragraph in table.rows[1].cells[2].paragraphs]
    detail_run_prototypes = [
        deepcopy(run._r)
        for run in table.rows[1].cells[5].paragraphs[0].runs
    ]
    for row_element in list(table._tbl.tr_lst[1:]):
        table._tbl.remove(row_element)

    pairs = _paired_photo_rows(data)
    if not pairs:
        pairs = [{"visible": None, "thermal": None}]

    for index, pair in enumerate(pairs, start=1):
        table._tbl.append(deepcopy(prototype_row))
        row = table.rows[-1]
        _set_paragraph_text(row.cells[0].paragraphs[0], str(index))
        _set_photo_cell(
            row.cells[1],
            pair["visible"],
            read_object=read_object,
            populated_prototypes=populated_photo_prototypes,
            empty_prototypes=empty_photo_prototypes,
        )
        _set_photo_cell(
            row.cells[2],
            pair["thermal"],
            read_object=read_object,
            populated_prototypes=populated_photo_prototypes,
            empty_prototypes=empty_photo_prototypes,
        )

        defects = [
            *(pair["visible"]["defects"] if pair["visible"] else []),
            *(pair["thermal"]["defects"] if pair["thermal"] else []),
        ]
        primary = pair["visible"] or pair["thermal"]
        _set_paragraph_text(
            row.cells[3].paragraphs[0],
            _metadata_text(primary["photo"] if primary else None),
        )
        _set_paragraph_text(row.cells[4].paragraphs[0], _defect_description(defects), blue=True)
        _set_defect_details(
            row.cells[5].paragraphs[0],
            defects,
            detail_run_prototypes,
        )

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_report_docx(
    report_title: str,
    report_no: str,
    report_data: dict[str, Any] | None,
    *,
    read_object: ObjectReader | None = None,
) -> bytes:
    if not FORMAL_TEMPLATE.is_file():
        raise DocxReportExportError("DOCX 导出模板不存在。")

    data = deepcopy(report_data or {})
    data["defects"] = number_defects(
        defect for defect in data.get("defects") or [] if isinstance(defect, dict)
    )
    document = Document(FORMAL_TEMPLATE)
    _set_title(document, data, report_title, report_no)
    _populate_result_table(document, data, read_object)
    _remove_trailing_empty_paragraph(document)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
