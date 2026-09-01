from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, TiffImagePlugin

from app import models  # noqa: F401
from app.db.base import Base
from app.services.docx_report import build_report_docx
from app.services.photo_metadata import (
    extract_formal_photo_metadata_from_bytes,
    extract_photo_metadata_from_bytes,
    facade_orientation_from_yaw,
    infer_drone_type,
)


def _dji_xmp(
    *,
    relative_altitude: str,
    gimbal_yaw_degree: str,
    camera_model: str | None = None,
    projection_pose: bool = False,
) -> bytes:
    model_attribute = f'tiff:Model="{camera_model}" '.encode() if camera_model else b""
    projection_attributes = (
        b'drone-dji:GpsLongitude="114.106574821" '
        b'drone-dji:GpsLatitude="22.578338627" '
        b'drone-dji:AbsoluteAltitude="36.533" '
        b'drone-dji:GimbalPitchDegree="-8.20" '
        b'drone-dji:GimbalRollDegree="0.00" '
        b'drone-dji:CalibratedFocalLength="3725.151611" '
        if projection_pose
        else b""
    )
    return (
        b"\xff\xd8"
        b'<rdf:Description drone-dji:ImageSource="Visible" '
        + model_attribute
        + projection_attributes
        + f'drone-dji:RelativeAltitude="{relative_altitude}" '.encode()
        + f'drone-dji:GimbalYawDegree="{gimbal_yaw_degree}" />'.encode()
        + b"\xff\xd9"
    )


def test_formal_photo_metadata_extracts_altitude_and_facade_orientation() -> None:
    data = _dji_xmp(
        relative_altitude="+42.500",
        gimbal_yaw_degree="90.0",
        camera_model="M3T",
        projection_pose=True,
    )

    metadata = extract_formal_photo_metadata_from_bytes(data)

    assert metadata["relative_altitude"] == 42.5
    assert metadata["gimbal_yaw_degree"] == 90.0
    assert metadata["longitude"] == pytest.approx(114.106574821)
    assert metadata["latitude"] == pytest.approx(22.578338627)
    assert metadata["absolute_altitude"] == 36.533
    assert metadata["gimbal_pitch_degree"] == -8.2
    assert metadata["gimbal_roll_degree"] == 0.0
    assert metadata["calibrated_focal_length"] == pytest.approx(3725.151611)
    assert metadata["facade_orientation"] == "西立面"
    assert metadata["camera_model"] == "M3T"
    assert metadata["drone_metadata_available"] is True
    assert metadata["professional_drone_photo"] is True


@pytest.mark.parametrize(
    ("data", "expected_reason"),
    [
        (
            _dji_xmp(relative_altitude="18.0", gimbal_yaw_degree="-45.0"),
            "missing-model",
        ),
        (b'\xff\xd8<rdf:Description tiff:Model="M3T" />\xff\xd9', "missing-drone-metadata"),
    ],
)
def test_formal_photo_requires_drone_metadata_and_camera_model(
    data: bytes,
    expected_reason: str,
) -> None:
    metadata = extract_formal_photo_metadata_from_bytes(data)

    assert metadata["professional_drone_photo"] is False, expected_reason


def test_formal_photo_reads_camera_model_from_exif() -> None:
    output = BytesIO()
    exif = Image.Exif()
    exif[0x010F] = "DJI"
    exif[0x0110] = "FC3582"
    Image.new("RGB", (8, 8), "white").save(output, format="JPEG", exif=exif)
    data = output.getvalue() + b'<rdf:Description drone-dji:RelativeAltitude="21.5" />'

    metadata = extract_formal_photo_metadata_from_bytes(data)

    assert metadata["camera_make"] == "DJI"
    assert metadata["camera_model"] == "FC3582"
    assert metadata["professional_drone_photo"] is True


def test_formal_photo_extracts_measurement_and_product_metadata() -> None:
    output = BytesIO()
    exif = Image.Exif()
    exif[0x010F] = "DJI"
    exif[0x0110] = "M4T"
    exif[0x920A] = TiffImagePlugin.IFDRational(6.72)
    exif[0xA405] = 24
    Image.new("RGB", (8, 8), "white").save(output, format="JPEG", exif=exif)
    data = output.getvalue() + (
        b'<rdf:Description drone-dji:RelativeAltitude="7.4" '
        b'drone-dji:ImageSource="WideCamera" '
        b'drone-dji:ProductName="DJI Matrice 4T" '
        b'drone-dji:DroneModel="M4T" '
        b'drone-dji:LRFTargetDistance="6.182" />'
    )

    metadata = extract_formal_photo_metadata_from_bytes(data)

    assert metadata["camera_product_name"] == "DJI Matrice 4T"
    assert metadata["drone_model"] == "M4T"
    assert metadata["focal_length_mm"] == pytest.approx(6.72)
    assert metadata["focal_length_35mm"] == 24
    assert metadata["lrf_target_distance"] == pytest.approx(6.182)
    assert infer_drone_type(metadata) == "dji_matrice_4t"


def test_formal_photo_accepts_non_dji_drone_metadata_namespace() -> None:
    data = (
        b"\xff\xd8"
        b'<rdf:Description tiff:Model="EVO II Pro" '
        b'drone-autel:RelativeAltitude="28.0" />'
        b"\xff\xd9"
    )

    metadata = extract_formal_photo_metadata_from_bytes(data)

    assert metadata["camera_model"] == "EVO II Pro"
    assert metadata["relative_altitude"] == 28.0
    assert metadata["professional_drone_photo"] is True


def test_trial_metadata_contract_does_not_include_formal_pose_fields() -> None:
    data = _dji_xmp(relative_altitude="18.0", gimbal_yaw_degree="-45.0")

    metadata = extract_photo_metadata_from_bytes(data)

    assert set(metadata) == {
        "xmp_drone_dji_image_source",
        "ifd0_image_description",
        "thermal_imaging_available",
    }


@pytest.mark.parametrize(
    ("yaw", "expected"),
    [
        (0.0, "南立面"),
        (45.0, "西南立面"),
        (90.0, "西立面"),
        (135.0, "西北立面"),
        (180.0, "北立面"),
        (225.0, "东北立面"),
        (270.0, "东立面"),
        (-45.0, "东南立面"),
        (None, None),
    ],
)
def test_facade_orientation_uses_reference_eight_direction_mapping(
    yaw: float | None,
    expected: str | None,
) -> None:
    assert facade_orientation_from_yaw(yaw) == expected


def test_formal_photo_table_persists_pose_metadata() -> None:
    photo = Base.metadata.tables["photo"]

    assert "relative_altitude" in photo.c
    assert "gimbal_yaw_degree" in photo.c
    assert "absolute_altitude" in photo.c
    assert "gimbal_pitch_degree" in photo.c
    assert "gimbal_roll_degree" in photo.c
    assert "calibrated_focal_length" in photo.c
    assert "focal_length_mm" in photo.c
    assert "focal_length_35mm" in photo.c
    assert "lrf_target_distance" in photo.c
    assert photo.c.relative_altitude.nullable
    assert photo.c.gimbal_yaw_degree.nullable


def test_formal_docx_adds_facade_orientation_and_capture_height_column() -> None:
    content = build_report_docx(
        "正式检测报告",
        "RPT-POSE-001",
        {
            "project": {"name": "13号楼"},
            "photos": [
                {
                    "id": "photo-1",
                    "facade_orientation": "西立面",
                    "relative_altitude": "42.500",
                }
            ],
            "defects": [
                {
                    "photo_id": "photo-1",
                    "photo_filename": "DJI_0001.JPG",
                    "defect_type": "crack",
                    "status": "confirmed",
                    "bbox_json": {},
                }
            ],
        },
    )

    with ZipFile(BytesIO(content)) as package:
        document = package.read("word/document.xml").decode("utf-8")

    assert "立面朝向" in document
    assert "拍摄高度" in document
    assert "西立面" in document
    assert "42.5 m" in document


def _report_jpeg() -> bytes:
    output = BytesIO()
    Image.new("RGB", (1600, 1200), "#D8C7A7").save(output, format="JPEG", quality=95)
    return output.getvalue()


def _has_red_box(image: Image.Image) -> bool:
    return any(
        red > 150 and red > green * 1.5 and red > blue * 1.3
        for red, green, blue in image.convert("RGB").getdata()
    )


def _has_blue_box(image: Image.Image) -> bool:
    return any(
        blue > 140 and blue > red * 1.5 and blue > green * 1.2
        for red, green, blue in image.convert("RGB").getdata()
    )


def _has_white_label_text(image: Image.Image) -> bool:
    return sum(
        1
        for red, green, blue in image.convert("RGB").getdata()
        if red > 235 and green > 235 and blue > 235
    ) > 20


def test_formal_docx_uses_reference_layout_and_pairs_four_by_three_photos() -> None:
    image_bytes = _report_jpeg()
    reads: list[tuple[str, str]] = []

    def read_object(bucket: str, object_key: str) -> bytes:
        reads.append((bucket, object_key))
        return image_bytes

    content = build_report_docx(
        "13号楼外立面表观病害筛查简报",
        "RPT-LAYOUT-001",
        {
            "project": {"name": "13号楼"},
            "photos": [
                {
                    "id": "visible-1",
                    "original_filename": "DJI_0165_V.JPG",
                    "storage_bucket": "inspection",
                    "storage_object_key": "photos/0165-v.jpg",
                    "photo_type": "visible",
                    "facade_orientation": "西立面",
                    "relative_altitude": "38.6",
                },
                {
                    "id": "thermal-1",
                    "original_filename": "DJI_0165_T.JPG",
                    "storage_bucket": "inspection",
                    "storage_object_key": "photos/0165-t.jpg",
                    "photo_type": "thermal",
                    "facade_orientation": "西立面",
                    "relative_altitude": "38.6",
                },
            ],
            "defects": [
                {
                    "photo_id": "visible-1",
                    "defect_type": "crack",
                    "bbox_json": {"x": 0.2, "y": 0.25, "width": 0.3, "height": 0.3},
                    "area": "0.777",
                    "area_estimated": True,
                    "length": "0.248",
                    "length_estimated": True,
                },
                {
                    "photo_id": "thermal-1",
                    "defect_type": "hollow",
                    "bbox_json": {"x": 800, "y": 300, "width": 480, "height": 360},
                    "area": "1.2",
                    "area_estimated": False,
                },
                {
                    "photo_id": "thermal-1",
                    "defect_type": "hollow",
                    "bbox_json": {"x": 0.1, "y": 0.6, "width": 0.2, "height": 0.2},
                },
            ],
        },
        read_object=read_object,
    )

    document = Document(BytesIO(content))
    assert document.sections[0].page_width.inches == pytest.approx(8.5)
    assert document.sections[0].left_margin.inches == pytest.approx(1.25)
    assert document.paragraphs[1].text == "13号楼-外立面表观病害筛查简报"
    assert document.paragraphs[1].runs[0].font.name == "黑体"
    assert document.paragraphs[1].runs[0].font.size.pt == pytest.approx(16)
    assert document.paragraphs[2].text == (
        "经对巡检结果进行空间定位与尺度估算，得到以下疑似病害位置。"
        "深度估计结果存在模型与相机参数误差，建议结合现场复核。"
    )
    assert "与面积" not in document.paragraphs[2].text

    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == [
        "序号",
        "可见光图像",
        "热红外图像",
        "立面朝向\n拍摄高度",
        "说明",
        "缺陷详情",
    ]
    assert [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.gridCol_lst
    ] == [526, 2548, 2546, 1090, 1864, 1459]
    assert len(table.rows) == 2
    assert all(
        cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for row in table.rows
        for cell in row.cells
    )
    assert table.rows[1].cells[1].text.endswith("DJI_0165_V.JPG")
    assert table.rows[1].cells[2].text.endswith("DJI_0165_T.JPG")
    assert table.rows[1].cells[3].text == "西立面\n38.6 m"
    assert table.rows[1].cells[4].text == "疑似裂缝: 1处\n疑似空鼓: 2处"
    assert table.rows[1].cells[5].text == (
        "裂缝-001≈0.248m\n"
        "空鼓-001 1.200m²\n"
        "参数不足"
    )
    assert "0.777" not in table.rows[1].cells[5].text
    detail_runs = table.rows[1].cells[5].paragraphs[0].runs
    assert detail_runs[3].text == "m\n"
    assert detail_runs[3].font.name == "微软雅黑"
    assert detail_runs[3].font.size.pt == pytest.approx(9)
    assert str(detail_runs[3].font.color.rgb) == "475467"
    assert reads == [
        ("inspection", "photos/0165-v.jpg"),
        ("inspection", "photos/0165-t.jpg"),
    ]

    assert len(document.inline_shapes) == 3
    logo, *table_images = document.inline_shapes
    assert logo.width == 1_463_040
    assert logo.height == 448_310
    embedded_images = []
    for shape in table_images:
        assert shape.width == 1_440_000
        assert shape.height == 1_080_000
        relationship_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = document.part.related_parts[relationship_id]
        with Image.open(BytesIO(image_part.blob)) as image:
            assert image.format == "JPEG"
            assert image.size == (1200, 900)
            embedded_images.append(image.copy())

    assert _has_red_box(embedded_images[0])
    assert _has_blue_box(embedded_images[1])
    assert all(_has_white_label_text(image) for image in embedded_images)


def test_formal_docx_shows_one_parameter_warning_when_measurements_are_missing() -> None:
    content = build_report_docx(
        "参数不足报告",
        "RPT-NO-MEASUREMENTS",
        {
            "project": {"name": "参数不足项目"},
            "photos": [
                {
                    "id": "visible-1",
                    "original_filename": "facade.jpg",
                    "photo_type": "visible",
                    "facade_orientation": "东立面",
                    "relative_altitude": "12.3",
                }
            ],
            "defects": [
                {
                    "photo_id": "visible-1",
                    "defect_type": "crack",
                    "bbox_json": {"x": 10, "y": 20, "width": 100, "height": 50},
                },
                {
                    "photo_id": "visible-1",
                    "defect_type": "hollow",
                    "bbox_json": {"x": 200, "y": 100, "width": 80, "height": 60},
                },
            ],
        },
    )

    document = Document(BytesIO(content))
    table = document.tables[0]
    assert all(len(row._tr.tc_lst) == 6 for row in table.rows)
    assert table.rows[0].cells[5].text == "缺陷详情"
    assert table.rows[1].cells[5].text == "参数不足"
